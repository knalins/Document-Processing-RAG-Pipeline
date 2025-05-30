"""
Document Processing Pipeline with LLMWhisperer, RAG Engine, and PDF Conversion
================================================================================
This script processes PDF documents, extracts questions, generates answers using RAG,
and converts the results back to PDF format.
"""

# ========================================
# IMPORTS
# ========================================

# Data Extraction by LLMWhisperer
from unstract.llmwhisperer import LLMWhispererClientV2
from unstract.llmwhisperer.client_v2 import LLMWhispererClientException
import time

# RAG Engine
import google.generativeai as genai
from llama_index.core import SimpleDirectoryReader
from llama_index.core import VectorStoreIndex
from IPython.display import Markdown, display
from llama_index.core import ServiceContext  # Delete it
from llama_index.core import Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import StorageContext, load_index_from_storage
from llama_index.embeddings.google_genai import GoogleGenAIEmbedding
from llama_index.llms.google_genai import GoogleGenAI
from google.genai.types import EmbedContentConfig

# Utilities
import os
import re
import logging
from datetime import datetime
from dotenv import load_dotenv

# Adobe PDF Services
from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
from adobe.pdfservices.operation.exception.exceptions import ServiceApiException, ServiceUsageException, SdkException
from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
from adobe.pdfservices.operation.io.stream_asset import StreamAsset
from adobe.pdfservices.operation.pdf_services import PDFServices
from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
from adobe.pdfservices.operation.pdfjobs.jobs.create_pdf_job import CreatePDFJob
from adobe.pdfservices.operation.pdfjobs.result.create_pdf_result import CreatePDFResult

# String to Docx
from docx import Document
from docx.shared import Pt

# Load environment variables
load_dotenv()

# =============================================================================
# CORE FUNCTIONS
# =============================================================================

def API_Setup_Files_Load(WhispererAPIKEY, GoogleAPIKEY, Localdatabase_Address, Input_File_Path):
    """
    Setup APIs and process document extraction

    Args:
        WhispererAPIKEY (str): API key for LLMWhisperer
        GoogleAPIKEY (str): API key for Google GenAI
        Localdatabase_Address (str): Path to local database file
        Input_File_Path (str): Path to input file for processing

    Returns:
        tuple: (extraction_result, vector_index)
    """
    # LLMWhisperer Setup
    client = LLMWhispererClientV2(
        base_url="https://llmwhisperer-api.us-central.unstract.com/api/v2",
        api_key=WhispererAPIKEY
    )

    # Google GenAI Gemini API Setup
    genai.configure(api_key=GoogleAPIKEY)
    Settings.llm = GoogleGenAI(models='gemini-1.5-flash-latest', api_key=GoogleAPIKEY)  # Use gemini-2.0-flash
    Settings.embed_model = GoogleGenAIEmbedding(
        model_name="models/embedding-001",
        api_key=GoogleAPIKEY
    )
    Settings.node_parser = SentenceSplitter(chunk_size=800, chunk_overlap=20)
    # Local database for references & retrieval
    documents = SimpleDirectoryReader(input_files=[Localdatabase_Address])
    doc = documents.load_data()
    index = VectorStoreIndex.from_documents(doc)

    # LLM Whisperer Working Data Extraction
    try:
        result = client.whisper(file_path=Input_File_Path)

        if result["status_code"] == 202:
            print("Whisper request accepted.")
            print(f"Whisper hash: {result['whisper_hash']}")

            while True:
                print("Polling for whisper status...")
                status = client.whisper_status(whisper_hash=result["whisper_hash"])

                if status["status"] == "processing":
                    print("STATUS: processing...")
                elif status["status"] == "delivered":
                    print("STATUS: Already delivered!")
                    break
                elif status["status"] == "unknown":
                    print("STATUS: unknown...")
                    break
                elif status["status"] == "processed":
                    print("STATUS: processed!")
                    print("Let's retrieve the result of the extraction...")
                    resulty = client.whisper_retrieve(whisper_hash=result["whisper_hash"])
                    print(resulty)
                    return resulty, index

                # Poll every 5 seconds
                time.sleep(5)

    except LLMWhispererClientException as e:
        print(e)


def extract_all_questions(document_text):
    """
    Extract different types of questions from document text

    Args:
        document_text (str): Raw extracted text from document

    Returns:
        dict: Dictionary containing numbered_questions, table_questions, and mcq_raw_sections
    """
    numbered_questions = []
    table_questions = []
    mcq_raw_sections = []

    # Normalize newlines and handle form feed character
    document_text = re.sub(r'\r\n', '\n', document_text)
    # Replace the form feed character with a consistent marker for splitting
    document_text = document_text.replace('<<<', '<<<\n---FORM_FEED---\n')
    # Reduce multiple blank lines to single blank lines
    document_text = re.sub(r'\n\s*\n', '\n\n', document_text).strip()

    # --- 1. Extract numbered questions ---
    numbered_questions_section_match = re.search(
        r'Questions for Test\s*\n(.*?)(?=\n\s*Question\s+Answer|\Z)',
        document_text,
        re.DOTALL
    )

    if numbered_questions_section_match:
        numbered_questions_text = numbered_questions_section_match.group(1).strip()
        questions_raw = re.split(r'^\s*\d+\.\s*', numbered_questions_text, flags=re.MULTILINE)

        for q_raw in questions_raw:
            cleaned_q = q_raw.strip()
            if cleaned_q:
                # Replace internal newlines and multiple spaces with a single space
                cleaned_q = re.sub(r'\s*\n\s*', ' ', cleaned_q)
                numbered_questions.append(cleaned_q)

    # --- 2. Extract table questions ---
    table_section_match = re.search(
        r'Question\s+Answer\s*\n(.*?)(?=\n<<<|\n\s*Tick all that applies:|\Z)',
        document_text,
        re.DOTALL
    )

    if table_section_match:
        table_content = table_section_match.group(1).strip()
        current_question = []

        for line in table_content.split('\n'):
            stripped_line = line.strip()
            if not stripped_line:
                continue  # Skip empty lines

            # If the line starts with a capital letter or a commonly used question word, it's likely a new question.
            # Otherwise, it's a continuation of the previous question.
            if re.match(r'^[A-Z][a-zA-Z]*', stripped_line) and not current_question:
                # First line of a new question
                current_question.append(stripped_line)
            elif re.match(r'^[A-Z][a-zA-Z]*', stripped_line) and current_question and len(current_question[0].split()) > 1:
                # New question, but the previous one was complete
                table_questions.append(" ".join(current_question))
                current_question = [stripped_line]
            elif current_question:
                # Continuation of the current question
                current_question.append(stripped_line)
            else:
                # Handle cases where the first line doesn't fit the 'new question' pattern perfectly
                current_question.append(stripped_line)

        # Add the last accumulated question
        if current_question:
            table_questions.append(" ".join(current_question))

    # --- 3. Extract raw MCQ section ---
    # This section is specifically between '<<<\n---FORM_FEED---\n' and the next '<<<\n---FORM_FEED---\n' or end of document
    mcq_section_match = re.search(
        r'Tick all that applies:.*?(?=\n---FORM_FEED---|\Z)',
        document_text,
        re.DOTALL
    )

    if mcq_section_match:
        mcq_raw_sections.append(mcq_section_match.group(0).strip())

    return {
        'numbered_questions': numbered_questions,
        'table_questions': table_questions,
        'mcq_raw_sections': mcq_raw_sections
    }


def clean_xml_incompatible_chars(text):
    """
    Removes characters that are not compatible with XML.
    This includes NULL bytes and other control characters (except common ones like tab, newline, carriage return).
    """
    control_chars_regex = r'[\x00-\x08\x0B\x0C\x0E-\x1F]'
    cleaned_text = re.sub(control_chars_regex, '', text)
    return cleaned_text

def convert_string_to_docx(input_string, filename="output.docx", font_name="Calibri", font_size=12):
    """
    Converts a given string into a DOCX file, cleaning incompatible XML characters first.

    Args:
        input_string (str): The string content to be written to the DOCX file.
        filename (str): The desired name for the output DOCX file.
        font_name (str): Optional. The name of the font to use.
        font_size (int): Optional. The font size in points.
    """
    # Clean the input string before processing
    cleaned_input_string = clean_xml_incompatible_chars(input_string)

    document = Document()

    # Add a paragraph and run to insert the string content
    paragraph = document.add_paragraph()
    run = paragraph.add_run(cleaned_input_string) # Use the cleaned string

    # Apply font and size
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)

    try:
        document.save(filename)
        print(f"Successfully converted string to '{filename}' with font '{font_name}' and size {font_size}.")
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")


class docxToPdfConverter:
    """
    Convert text files to PDF using Adobe PDF Services
    """

    def __init__(self, input_txt_path, output_pdf_path, client_id, client_secret):
        self.input_txt_path = input_txt_path
        self.output_pdf_path = output_pdf_path
        self.client_id = client_id
        self.client_secret = client_secret
        self.process_conversion()

    def process_conversion(self):
        """Process the text to PDF conversion"""
        try:
            # 1. Read the input docx file as a byte stream
            with open(self.input_txt_path, 'rb') as file:
                input_stream = file.read()

            # 2. Initial setup: create credentials instance using ServicePrincipalCredentials
            credentials = ServicePrincipalCredentials(
                client_id=self.client_id,
                client_secret=self.client_secret
            )

            # 3. Creates a PDF Services instance
            pdf_services = PDFServices(credentials=credentials)

            # 4. Creates an asset from the source file and uploads it to Adobe PDF Services
            input_asset = pdf_services.upload(input_stream=input_stream, mime_type=PDFServicesMediaType.DOCX)

            # 5. Creates a new CreatePDFJob instance
            create_pdf_job = CreatePDFJob(input_asset=input_asset)

            # 6. Submit the job and gets the job result
            logging.info(f"Attempting to convert '{self.input_txt_path}' to PDF...")
            location = pdf_services.submit(create_pdf_job)
            pdf_services_response = pdf_services.get_job_result(location, CreatePDFResult)

            # 7. Get content from the resulting asset
            result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)

            # 8. Create the output directory if it doesn't exist
            os.makedirs(os.path.dirname(self.output_pdf_path), exist_ok=True)

            # 9. Creates an output stream and copy stream asset's content to it
            with open(self.output_pdf_path, "wb") as file:
                file.write(stream_asset.get_input_stream())

            logging.info(f"Successfully converted '{self.input_txt_path}' to '{self.output_pdf_path}'")

        except (ServiceApiException, ServiceUsageException, SdkException) as e:
            logging.exception(f'Exception encountered while executing operation: {e}')
        except Exception as e:
            logging.exception(f'An unexpected error occurred: {e}')


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main execution function"""

    # Configure logging for better visibility into SDK operations
    logging.basicConfig(level=logging.INFO)

    # =============================================================================
    # STEP 1: API SETUP AND DATA EXTRACTION (Starter)
    # =============================================================================

    resulty, index = API_Setup_Files_Load(
        WhispererAPIKEY="WhispererLLM API Key",#Put Your WhispererLLM API Key Here
        GoogleAPIKEY="Google GenAI API Key", #Put Your Google GenAI API Key Here
        Localdatabase_Address="Dataformodel.txt", #Put Your Local Database Address for RAG Engine Here
        Input_File_Path="Questions_for_Test.pdf" #Put Your Input File Path Here
    )

    # Configuration
    CLIENT_ID = "PDF Services Client ID" #Put Your Adobe PDF Services Client ID from Adobe Developer Console Here
    CLIENT_SECRET = "PDF Services Client Secret" #Put Your Adobe PDF Services Client Secret from Adobe Developer Console Here
    # =============================================================================
    # STEP 2: EXTRACT AND ORGANIZE DATA
    # =============================================================================

    # Storing PDF Data in String format in location Variable(extracted_data)
    extracted_data = resulty['extraction']['result_text']
    print(type(extracted_data))  # String
    print(extracted_data)

    # Extraction of Questions for LLMs from Raw Text
    Organised_Extracted_data = extract_all_questions(extracted_data)
    print(Organised_Extracted_data)

    # =============================================================================
    # STEP 3: PREPARE QUERIES FOR OUTPUT ENGINE AND GENERATE RESPONSES
    # =============================================================================
    index.storage_context.persist()
    query_engine = index.as_query_engine()
    response_from_engine = ""
    for(key, value) in Organised_Extracted_data.items():
        data=Organised_Extracted_data[key]
        if(key=="numbered_questions"):
            engine_query = "Give both question and Answer\n" + '\n'.join(data)
        elif(key=="table_questions"):
            engine_query = "Keep both Question with answers in one line\n" + '\n'.join(data)
        elif(key=="mcq_raw_sections"):
            engine_query = "Keep same text with marked answers\n" + '\n'.join(data)
        response_from_engine = response_from_engine+query_engine.query(engine_query).response+"\n"


    # =============================================================================
    # STEP 5: SAVE RESULTS TO DOCX FILE
    # =============================================================================
    Local_Output_Path=f"Output{datetime.now().strftime('%Y-%m-%dT%H-%M-%S')}.docx"

    convert_string_to_docx(
        response_from_engine, # Use your 'extracted_data' variable here
        filename=Local_Output_Path, ## Output Filename
        font_name="Arial",
        font_size=14
        )

    # f = open(Local_Output_Path, "w")
    # f.write(response_from_engine)
    # f.close()
    # =============================================================================
    # STEP 6: CONVERT DOCX TO PDF
    # =============================================================================

    # Define your input and output file paths

    OUTPUT_PDF_DIRECTORY = "output/docxToPdfConversion"  # Directory for output
    OUTPUT_PDF_FILENAME = f"output_{Local_Output_Path}.pdf"
    INPUT_docx_FILE = Local_Output_Path ##File For pdf creation
    OUTPUT_PDF_PATH = os.path.join(OUTPUT_PDF_DIRECTORY, OUTPUT_PDF_FILENAME)
    # Run the conversion
    converter = docxToPdfConverter(INPUT_docx_FILE, OUTPUT_PDF_PATH, CLIENT_ID, CLIENT_SECRET)


# =============================================================================
# EXECUTION
# =============================================================================

if __name__ == "__main__":
    main()